import os
import re
from datetime import datetime

# --- IMPORTACIONES PDF (ReportLab) ---
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.lib.enums import TA_LEFT, TA_JUSTIFY, TA_CENTER
from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, Image, 
                                Table, TableStyle, PageBreak, ListFlowable, ListItem)
from reportlab.platypus.tableofcontents import TableOfContents

# --- IMPORTACIONES WORD (python-docx) ---
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# --- CONFIGURACIÓN GLOBAL ---
HEADER_LOGO_PATH = "CTM_Agents/static/CotecmarLogo.png"
COVER_IMAGE_PATH = "CTM_Agents/generated_images/detección_de_anomalías_sísmicas_con_inteligencia_artificial.png"
OUTPUT_DIR = "test_reports_final"

# =============================================================================================
#                                  GENERADOR DE PDF (CLASE)
# =============================================================================================

class MecanismoTOC(SimpleDocTemplate):
    def afterFlowable(self, flowable):
        """Detecta títulos para el índice del PDF."""
        if flowable.__class__.__name__ == 'Paragraph':
            text = flowable.getPlainText()
            style = flowable.style.name
            if style == 'CustomH1':
                self.notify('TOCEntry', (0, text, self.page))
            elif style == 'CustomH2':
                self.notify('TOCEntry', (1, text, self.page))

class ProjectPDFGenerator:
    # Colores específicos para ReportLab (Hex)
    COLOR_PRIMARY = colors.HexColor('#003366')
    COLOR_SECONDARY = colors.HexColor('#0066CC')
    COLOR_ACCENT = colors.HexColor('#E6F0FF')
    COLOR_TEXT = colors.HexColor('#2A2A2A')

    def __init__(self, filename, data):
        self.filename = filename
        self.data = data
        self.styles = self._create_styles()
        
    def _create_styles(self):
        styles = getSampleStyleSheet()
        # Estilos PDF
        styles.add(ParagraphStyle(name='CustomH1', parent=styles['Heading1'], fontSize=18, textColor=self.COLOR_PRIMARY, spaceAfter=12, spaceBefore=12, fontName='Helvetica-Bold', borderPadding=0))
        styles.add(ParagraphStyle(name='CustomH2', parent=styles['Heading2'], fontSize=14, textColor=self.COLOR_SECONDARY, spaceAfter=10, spaceBefore=12, fontName='Helvetica-Bold'))
        styles.add(ParagraphStyle(name='CustomH3', parent=styles['Heading3'], fontSize=12, textColor=colors.HexColor('#444444'), spaceAfter=8, spaceBefore=10, fontName='Helvetica-BoldOblique'))
        styles.add(ParagraphStyle(name='CustomBody', parent=styles['Normal'], fontSize=10, leading=14, alignment=TA_JUSTIFY, textColor=self.COLOR_TEXT, spaceAfter=6))
        
        # Estilo APA (Sangría Francesa)
        styles.add(ParagraphStyle(name='APA_Reference', parent=styles['Normal'], fontSize=10, leading=14, alignment=TA_LEFT, leftIndent=36, firstLineIndent=-36, spaceAfter=6))
        
        # Tablas y TOC
        styles.add(ParagraphStyle(name='CellHeader', parent=styles['Normal'], fontSize=9, fontName='Helvetica-Bold', textColor=colors.white, alignment=TA_CENTER))
        styles.add(ParagraphStyle(name='CellBody', parent=styles['Normal'], fontSize=9, leading=11, alignment=TA_LEFT))
        styles.add(ParagraphStyle(name='TOCHeading', parent=styles['Heading1'], fontSize=16, textColor=self.COLOR_PRIMARY, spaceAfter=10, alignment=TA_CENTER))
        styles.add(ParagraphStyle(name='TOCEntry1', parent=styles['Normal'], fontSize=11, fontName='Helvetica-Bold', spaceAfter=5, textColor=self.COLOR_PRIMARY))
        styles.add(ParagraphStyle(name='TOCEntry2', parent=styles['Normal'], fontSize=10, leftIndent=20, spaceAfter=2, textColor=self.COLOR_TEXT))
        return styles

    def _header_footer_layout(self, canvas, doc):
        canvas.saveState()
        page_width, page_height = letter
        margin = 2 * cm
        # Header
        logo_w, logo_h = 4 * cm, 1.5 * cm
        y_pos = page_height - 2.5 * cm
        x_pos = page_width - margin - logo_w
        if os.path.exists(HEADER_LOGO_PATH):
            try: canvas.drawImage(HEADER_LOGO_PATH, x_pos, y_pos, width=logo_w, height=logo_h, mask='auto', preserveAspectRatio=True)
            except: pass
        else:
            canvas.setFillColor(colors.lightgrey); canvas.rect(x_pos, y_pos, logo_w, logo_h, fill=1)
        # Footer
        canvas.setStrokeColor(self.COLOR_SECONDARY); canvas.setLineWidth(0.5)
        canvas.line(margin, 2*cm, page_width - margin, 2*cm)
        canvas.setFont('Helvetica', 8); canvas.setFillColor(colors.gray)
        canvas.drawString(margin, 1.5*cm, f"Proyecto: {self.data['titulo_corto']}")
        canvas.drawRightString(page_width - margin, 1.5*cm, f"Página {doc.page}")
        canvas.restoreState()

    def create_table(self, data, col_widths=None):
        formatted_data = [[Paragraph(c, self.styles['CellHeader'] if i==0 else self.styles['CellBody']) if isinstance(c, str) else c for c in r] for i, r in enumerate(data)]
        t = Table(formatted_data, colWidths=col_widths)
        style_cmds = [('BACKGROUND', (0,0), (-1,0), self.COLOR_SECONDARY), ('TEXTCOLOR', (0,0), (-1,0), colors.white), ('ALIGN', (0,0), (-1,-1), 'LEFT'), ('VALIGN', (0,0), (-1,-1), 'MIDDLE'), ('GRID', (0,0), (-1,-1), 0.5, colors.grey), ('PADDING', (0,0), (-1,-1), 6)]
        for i in range(1, len(data)): style_cmds.append(('BACKGROUND', (0,i), (-1,i), self.COLOR_ACCENT if i%2==0 else colors.white))
        t.setStyle(TableStyle(style_cmds))
        return t

    def generate(self):
        doc = MecanismoTOC(self.filename, pagesize=letter, rightMargin=2*cm, leftMargin=2*cm, topMargin=3*cm, bottomMargin=2.5*cm)
        story = []
        S = self.styles
        
        # Contenido PDF
        story.append(Paragraph("Estructura del Proyecto", S['CustomH1'])); story.append(Spacer(1, 1*cm))
        story.append(Paragraph("Tabla de Contenido", S['TOCHeading']))
        toc = TableOfContents(); toc.levelStyles = [S['TOCEntry1'], S['TOCEntry2']]; story.append(toc); story.append(PageBreak())

        # Secciones
        story.append(Paragraph("Generalidades", S['CustomH1']))
        story.append(Paragraph("Información Básica", S['CustomH2']))
        story.append(self.create_table([["Atributo", "Valor"], ["Código", self.data['codigo']], ["Título", self.data['titulo_completo']], ["Entidad", self.data['entidad']]], col_widths=[6*cm, 11*cm]))
        
        story.append(Paragraph("Resumen Ejecutivo", S['CustomH1']))
        story.append(Paragraph(self.data['resumen_ejecutivo'], S['CustomBody']))
        
        story.append(Paragraph("Referencias (APA)", S['CustomH1']))
        for ref in self.data['referencias']: story.append(Paragraph(ref, S['APA_Reference']))

        doc.multiBuild(story, onFirstPage=self._header_footer_layout, onLaterPages=self._header_footer_layout)
        print(f"✅ [PDF] Generado: {self.filename}")

# =============================================================================================
#                                  GENERADOR DE WORD (CLASE)
# =============================================================================================

class ProjectWordGenerator:
    # Colores específicos para Word (RGB)
    COLOR_PRIMARY = RGBColor(0, 51, 102)
    COLOR_SECONDARY = RGBColor(0, 102, 204)
    
    def __init__(self, filename, data):
        self.filename = filename
        self.data = data
        self.doc = Document()
        self._setup_styles()
        self._setup_header_footer()

    def _setup_styles(self):
        styles = self.doc.styles
        # H1
        h1 = styles['Heading 1']; h1.font.name = 'Arial'; h1.font.size = Pt(16); h1.font.color.rgb = self.COLOR_PRIMARY; h1.font.bold = True
        # H2
        h2 = styles['Heading 2']; h2.font.name = 'Arial'; h2.font.size = Pt(14); h2.font.color.rgb = self.COLOR_SECONDARY; h2.font.bold = True
        # APA
        if 'APA_Reference' not in styles:
            apa = styles.add_style('APA_Reference', 1)
            apa.base_style = styles['Normal']; apa.font.name = 'Arial'; apa.font.size = Pt(10)
            apa.paragraph_format.left_indent = Cm(1.27); apa.paragraph_format.first_line_indent = Cm(-1.27)

    def _setup_header_footer(self):
        section = self.doc.sections[0]
        section.header_distance = Cm(1.0)
        # Header Imagen
        header_para = section.header.paragraphs[0]; header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if os.path.exists(HEADER_LOGO_PATH):
            try: run = header_para.add_run(); run.add_picture(HEADER_LOGO_PATH, width=Cm(4))
            except: pass
        # Footer Text
        footer_para = section.footer.paragraphs[0]; footer_para.text = f"Proyecto: {self.data['titulo_corto']} | Generado Automáticamente"
        footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    def add_html_paragraph(self, text, style='Normal'):
        p = self.doc.add_paragraph(style=style)
        tokens = re.split(r'(<b>.*?</b>|<i>.*?</i>)', text)
        for token in tokens:
            if not token: continue
            run = p.add_run()
            if token.startswith('<b>'): run.text = token[3:-4]; run.bold = True
            elif token.startswith('<i>'): run.text = token[3:-4]; run.italic = True
            else: run.text = token

    def create_table(self, data):
        table = self.doc.add_table(rows=len(data), cols=len(data[0])); table.style = 'Table Grid'
        for i, row in enumerate(data):
            for j, val in enumerate(row):
                cell = table.rows[i].cells[j]; cell.text = str(val)
                if i == 0: 
                    for p in cell.paragraphs: 
                        for r in p.runs: r.font.bold = True; r.font.color.rgb = self.COLOR_PRIMARY

    def add_toc_xml(self):
        p = self.doc.add_paragraph()
        run = p.add_run(); 
        fldChar = OxmlElement('w:fldChar'); fldChar.set(qn('w:fldCharType'), 'begin'); run._r.append(fldChar)
        instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve'); instr.text = 'TOC \\o "1-3" \\h \\z \\u'; run._r.append(instr)
        fldChar = OxmlElement('w:fldChar'); fldChar.set(qn('w:fldCharType'), 'separate'); run._r.append(fldChar)
        run = p.add_run("Clic derecho -> Actualizar campos para ver Tabla de Contenido"); run.font.italic = True
        fldChar = OxmlElement('w:fldChar'); fldChar.set(qn('w:fldCharType'), 'end'); run._r.append(fldChar)

    def generate(self):
        self.doc.add_heading('Estructura del Proyecto', 0)
        self.doc.add_heading('Tabla de Contenido', 1)
        self.add_toc_xml()
        self.doc.add_page_break()

        self.doc.add_heading('Generalidades', 1)
        self.doc.add_heading('Información Básica', 2)
        self.create_table([["Atributo", "Valor"], ["Código", self.data['codigo']], ["Título", self.data['titulo_completo']], ["Entidad", self.data['entidad']]])
        
        self.doc.add_heading('Resumen Ejecutivo', 1)
        self.doc.add_paragraph(self.data['resumen_ejecutivo'])

        self.doc.add_heading('Referencias (APA)', 1)
        for ref in self.data['referencias']: self.add_html_paragraph(ref, style='APA_Reference')

        self.doc.save(self.filename)
        print(f"✅ [WORD] Generado: {self.filename}")

# =============================================================================================
#                                  EJECUCIÓN PRINCIPAL
# =============================================================================================

DUMMY_DATA = {
    "titulo_corto": "Sismología IA",
    "titulo_completo": "Detección de Anomalías Sísmicas con Inteligencia Artificial",
    "codigo": "109755-2025",
    "convocatoria": "COLOMBIA INTELIGENTE",
    "entidad": "ESCUELA NAVAL DE CADETES",
    "keywords": ["IA", "Sismología", "IoT"],
    "resumen_ejecutivo": "Este proyecto busca democratizar el acceso a sistemas de alerta temprana mediante sensores IoT...",
    "entidades": [["Escuela Naval", "Ejecutor"], ["Univ. Cartagena", "Aliado"]],
    "referencias": [
        "Airlangga, G. (2024). ML Techniques. <i>Journal of Geophysics, 45</i>(2), 112-130.",
        "Lin, J. (2025). Anomaly detection. <i>Seismological Research Letters, 96</i>(1), 45-58."
    ]
}

if __name__ == "__main__":
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    timestamp = datetime.now().strftime('%Y%m%d_%H%M')
    
    # 1. GENERAR PDF
    pdf_filename = f"Proyecto_{timestamp}.pdf"
    pdf_path = os.path.join(OUTPUT_DIR, pdf_filename)
    pdf_gen = ProjectPDFGenerator(pdf_path, DUMMY_DATA)
    pdf_gen.generate()
    
    # 2. GENERAR WORD
    word_filename = f"Proyecto_{timestamp}.docx"
    word_path = os.path.join(OUTPUT_DIR, word_filename)
    word_gen = ProjectWordGenerator(word_path, DUMMY_DATA)
    word_gen.generate()

    print("\n🚀 PROCESO FINALIZADO: Ambos documentos han sido creados.")