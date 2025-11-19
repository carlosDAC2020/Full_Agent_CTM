import os
import re
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def convert_readme_to_word(readme_path, output_path):
    """
    Converts a specific README.md file to a Word document, handling headers,
    lists, and images.
    """
    if not os.path.exists(readme_path):
        print(f"Error: File not found at {readme_path}")
        return

    doc = Document()
    
    # Set default font to something standard if desired, though defaults are usually fine
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    with open(readme_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    # Simple state machine / parser
    in_code_block = False
    
    for line in lines:
        line = line.strip()
        
        # Skip empty lines mostly, but maybe keep them as paragraph breaks if needed
        # For this simple parser, we might just skip completely empty lines unless we want spacing
        if not line:
            continue

        # Handle Headers
        if line.startswith('#'):
            level = len(line.split()[0])
            text = line.lstrip('#').strip()
            # Word supports headings 1-9
            if level > 9: level = 9
            doc.add_heading(text, level=level)
            continue

        # Handle Horizontal Rule
        if line.startswith('---') or line.startswith('***'):
            doc.add_paragraph('_' * 50) # Simple visual separator
            continue

        # Handle Images
        # Regex for ![alt](path)
        img_match = re.match(r'!\[(.*?)\]\((.*?)\)', line)
        if img_match:
            alt_text = img_match.group(1)
            img_path = img_match.group(2)
            
            # Resolve relative paths
            if not os.path.isabs(img_path):
                base_dir = os.path.dirname(readme_path)
                img_full_path = os.path.join(base_dir, img_path)
            else:
                img_full_path = img_path
            
            if os.path.exists(img_full_path):
                try:
                    doc.add_picture(img_full_path, width=Inches(6))
                    # Add caption/alt text if desired
                    # last_paragraph = doc.paragraphs[-1] 
                    # last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                except Exception as e:
                    doc.add_paragraph(f"[Error inserting image: {e}]")
            else:
                doc.add_paragraph(f"[Image not found: {img_path}]")
            continue

        # Handle Lists
        if line.startswith('* ') or line.startswith('- '):
            text = line[2:].strip()
            # Handle bolding inside list items manually or just let it be plain text for now
            # A full markdown parser is complex, we'll do simple bold replacement
            p = doc.add_paragraph(style='List Bullet')
            add_text_with_formatting(p, text)
            continue
        
        if re.match(r'^\d+\.\s', line):
            # Ordered list
            parts = line.split('.', 1)
            if len(parts) > 1:
                text = parts[1].strip()
                p = doc.add_paragraph(style='List Number')
                add_text_with_formatting(p, text)
                continue

        # Normal Paragraph
        p = doc.add_paragraph()
        add_text_with_formatting(p, line)

    doc.save(output_path)
    print(f"Successfully created {output_path}")

def add_text_with_formatting(paragraph, text):
    """
    Parses **bold** text and adds runs to the paragraph.
    """
    # Split by ** to find bold parts
    # This is a very basic parser. It assumes balanced **.
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)

if __name__ == "__main__":
    readme_file = r"c:\Users\LABSISTEMAS03\Documents\carlos\Full_Agent_CTM\CTM_Agents\README.md"
    output_file = r"c:\Users\LABSISTEMAS03\Documents\carlos\Full_Agent_CTM\CTM_Agents\Reporte_CTM_Agent.docx"
    convert_readme_to_word(readme_file, output_file)
