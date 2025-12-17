import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def convert_markdown_to_word(markdown_path, output_path):
    """
    Converts a Markdown file to a Word document, handling headers,
    lists, tables, code blocks, and images.
    """
    if not os.path.exists(markdown_path):
        print(f"Error: File not found at {markdown_path}")
        return

    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    with open(markdown_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    # Simple state machine / parser
    in_code_block = False
    in_table = False
    code_lines = []
    table_lines = []
    
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        
        # Handle code blocks
        if line.startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_lines = []
                i += 1
                continue
            else:
                # End of code block
                in_code_block = False
                if code_lines:
                    p = doc.add_paragraph()
                    p.style = 'Normal'
                    run = p.add_run('\n'.join(code_lines))
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                    # Add light gray background
                    shading_elm = OxmlElement('w:shd')
                    shading_elm.set(qn('w:fill'), 'F0F0F0')
                    p._element.get_or_add_pPr().append(shading_elm)
                code_lines = []
                i += 1
                continue
        
        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Handle tables
        if line.startswith('|') and '|' in line[1:]:
            if not in_table:
                in_table = True
                table_lines = [line]
            else:
                table_lines.append(line)
            i += 1
            # Check if next line is still part of table
            if i < len(lines) and not lines[i].strip().startswith('|'):
                # End of table
                in_table = False
                add_table_to_doc(doc, table_lines)
                table_lines = []
            continue
        
        # Skip empty lines mostly
        if not line.strip():
            # Add a small paragraph break for spacing
            if i > 0 and lines[i-1].strip():
                doc.add_paragraph()
            i += 1
            continue

        # Handle Headers
        if line.startswith('#'):
            level = 0
            for char in line:
                if char == '#':
                    level += 1
                else:
                    break
            text = line.lstrip('#').strip()
            # Word supports headings 1-9
            if level > 9:
                level = 9
            doc.add_heading(text, level=level)
            i += 1
            continue

        # Handle Horizontal Rule
        if line.startswith('---') or line.startswith('***'):
            doc.add_paragraph('_' * 70)
            i += 1
            continue

        # Handle Images
        # Regex for ![alt](path)
        img_match = re.match(r'!\[(.*?)\]\((.*?)\)', line)
        if img_match:
            alt_text = img_match.group(1)
            img_path = img_match.group(2)
            
            # Resolve relative paths
            if not os.path.isabs(img_path):
                base_dir = os.path.dirname(markdown_path)
                img_full_path = os.path.join(base_dir, img_path)
            else:
                img_full_path = img_path
            
            if os.path.exists(img_full_path):
                try:
                    doc.add_picture(img_full_path, width=Inches(6))
                    # Add caption if there's alt text
                    if alt_text:
                        p = doc.add_paragraph(alt_text)
                        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        p.runs[0].italic = True
                        p.runs[0].font.size = Pt(10)
                except Exception as e:
                    doc.add_paragraph(f"[Error inserting image: {e}]")
            else:
                doc.add_paragraph(f"[Image not found: {img_path}]")
            i += 1
            continue

        # Handle Lists (with indentation support)
        indent_level = 0
        stripped = line.lstrip()
        if stripped.startswith('* ') or stripped.startswith('- '):
            indent_level = (len(line) - len(stripped)) // 2
            text = stripped[2:].strip()
            p = doc.add_paragraph(style='List Bullet')
            add_text_with_formatting(p, text)
            # Adjust indentation
            if indent_level > 0:
                p.paragraph_format.left_indent = Inches(0.25 * indent_level)
            i += 1
            continue
        
        if re.match(r'^\s*\d+\.?\s', line):
            # Ordered list
            indent_level = (len(line) - len(stripped)) // 2
            match = re.match(r'^\s*\d+\.?\s+(.*)', line)
            if match:
                text = match.group(1).strip()
                p = doc.add_paragraph(style='List Number')
                add_text_with_formatting(p, text)
                # Adjust indentation
                if indent_level > 0:
                    p.paragraph_format.left_indent = Inches(0.25 * indent_level)
                i += 1
                continue

        # Normal Paragraph
        p = doc.add_paragraph()
        add_text_with_formatting(p, line)
        i += 1

    doc.save(output_path)
    print(f"Successfully created {output_path}")


def add_table_to_doc(doc, table_lines):
    """
    Converts markdown table lines to a Word table.
    """
    if len(table_lines) < 2:
        return
    
    # Parse table
    rows = []
    for line in table_lines:
        # Skip separator lines (e.g., |---|---|)
        if re.match(r'^\|[\s\-:|]+\|$', line):
            continue
        cells = [cell.strip() for cell in line.split('|')[1:-1]]
        if cells:
            rows.append(cells)
    
    if not rows:
        return
    
    # Create table
    num_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Light Grid Accent 1'
    
    # Fill table
    for i, row_data in enumerate(rows):
        for j, cell_data in enumerate(row_data):
            if j < len(table.rows[i].cells):
                cell = table.rows[i].cells[j]
                # Add text with formatting
                add_text_with_formatting(cell.paragraphs[0], cell_data)
                # Make header row bold
                if i == 0:
                    for run in cell.paragraphs[0].runs:
                        run.bold = True

def add_text_with_formatting(paragraph, text):
    """
    Parses **bold**, *italic*, `code`, and [links](url) from text and adds runs to the paragraph.
    """
    # Handle inline code first
    parts = re.split(r'(`[^`]+`)', text)
    for part in parts:
        if part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
            # Light gray background for inline code
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), 'F0F0F0')
            run._element.get_or_add_rPr().append(shading_elm)
        else:
            # Handle bold and italic
            add_bold_italic_text(paragraph, part)

def add_bold_italic_text(paragraph, text):
    """
    Parses **bold** and *italic* text.
    """
    # Handle bold first
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            inner_text = part[2:-2]
            # Check for italic inside bold
            italic_parts = re.split(r'(\*.*?\*)', inner_text)
            for ipart in italic_parts:
                if ipart.startswith('*') and ipart.endswith('*') and not ipart.startswith('**'):
                    run = paragraph.add_run(ipart[1:-1])
                    run.bold = True
                    run.italic = True
                else:
                    run = paragraph.add_run(ipart)
                    run.bold = True
        else:
            # Check for italic
            italic_parts = re.split(r'(\*[^*]+\*)', part)
            for ipart in italic_parts:
                if ipart.startswith('*') and ipart.endswith('*') and not ipart.startswith('**'):
                    run = paragraph.add_run(ipart[1:-1])
                    run.italic = True
                else:
                    # Handle links [text](url)
                    link_parts = re.split(r'(\[.*?\]\(.*?\))', ipart)
                    for lpart in link_parts:
                        link_match = re.match(r'\[(.*?)\]\((.*?)\)', lpart)
                        if link_match:
                            link_text = link_match.group(1)
                            link_url = link_match.group(2)
                            run = paragraph.add_run(link_text)
                            run.font.color.rgb = RGBColor(0, 0, 255)
                            run.underline = True
                        else:
                            paragraph.add_run(lpart)

if __name__ == "__main__":
    # Paths for entregable #1
    markdown_file = r"c:\Users\dan20\Documents\Full_Agent_CTM\entregables\entregable#1.md"
    output_file = r"c:\Users\dan20\Documents\Full_Agent_CTM\entregables\entregable#1.docx"
    
    convert_markdown_to_word(markdown_file, output_file)
